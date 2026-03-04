"""文档导出工具

将 Markdown 内容转换为符合公文格式的 Word 文档。
基于 python-docx，参考 GB/T 9704 公文格式标准。
"""

from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from loguru import logger

from src.config import WRITABLE_DATA_DIR

EXPORT_DIR = WRITABLE_DATA_DIR / "exports"


def export_markdown_to_word(
    content: str,
    title: str = "",
    doc_type: str = "general",
) -> Path:
    """将 Markdown 内容导出为 Word 文档

    Args:
        content: Markdown 格式的文本内容
        title: 文档标题（用于文件名）
        doc_type: 文档类型 general/notice/report/speech

    Returns:
        生成的 Word 文件路径
    """
    EXPORT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # 设置页面边距（GB/T 9704 标准：上3.7cm 下3.5cm 左2.8cm 右2.6cm）
    for section in doc.sections:
        section.top_margin = Cm(3.7)
        section.bottom_margin = Cm(3.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)

    # 设置默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "仿宋"
    font.size = Pt(16)  # 三号字 ≈ 16pt
    font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.line_spacing = Pt(28)  # 行距 28 磅
    style.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进 2 字符

    # 解析 Markdown 并逐段写入
    lines = content.strip().split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # 跳过空行
        if not line.strip():
            i += 1
            continue

        # 一级标题 → 公文标题（二号方正小标宋）
        if line.startswith("# "):
            text = line[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(20)
            run = p.add_run(text)
            run.font.size = Pt(22)  # 二号字
            run.font.bold = True
            run.font.name = "方正小标宋简体"
            i += 1
            continue

        # 二级标题 → 一级标题"一、"（三号黑体）
        if line.startswith("## "):
            text = line[3:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(10)
            run = p.add_run(text)
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.name = "黑体"
            i += 1
            continue

        # 三级标题 → 二级标题"（一）"（三号楷体）
        if line.startswith("### "):
            text = line[4:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(6)
            run = p.add_run(text)
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.name = "楷体"
            i += 1
            continue

        # 分割线 → 跳过
        if line.strip().startswith("---"):
            i += 1
            continue

        # 普通段落
        text = _clean_markdown(line.strip())
        if text:
            p = doc.add_paragraph()
            _add_formatted_text(p, text)

        i += 1

    # 生成文件名
    safe_title = re.sub(r'[\\/:*?"<>|]', "_", title or "文档")[:50]
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{safe_title}_{timestamp}.docx"
    file_path = EXPORT_DIR / filename

    doc.save(str(file_path))
    logger.info(f"Word 文档已导出: {file_path}")
    return file_path


def _clean_markdown(text: str) -> str:
    """清理 Markdown 标记，保留纯文本"""
    # 粗体
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    # 斜体
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    # 行内代码
    text = re.sub(r"`(.*?)`", r"\1", text)
    # 列表标记
    text = re.sub(r"^[\-\*]\s+", "", text)
    text = re.sub(r"^\d+\.\s+", "", text)
    return text.strip()


def _add_formatted_text(paragraph, text: str):
    """向段落添加文本，处理粗体标记"""
    # 简单处理：直接添加清理后的文本
    run = paragraph.add_run(text)
    run.font.name = "仿宋"
    run.font.size = Pt(16)
